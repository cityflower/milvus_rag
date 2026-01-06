# milvus_project/utils/file_reader.py
import os
import fitz # PyMuPDF
from docx import Document
from typing import Literal

def read_file(path: str) -> str:
    """根据文件扩展名读取文件内容"""
    ext = os.path.splitext(path)[1].lower()
    
    if ext == ".pdf":
        # 使用 PyMuPDF (fitz) 读取 PDF
        with fitz.open(path) as doc:
            return "".join(page.get_text() for page in doc)
            
    elif ext in (".docx", ".doc"):
        # 使用 python-docx 读取 Word
        # 注意: .doc 文件需要额外的库支持，这里默认 Document(path) 可以处理
        try:
            return "\n".join(p.text for p in Document(path).paragraphs)
        except Exception as e:
            raise IOError(f"读取 Word 文件失败: {e}")
            
    else:
        # 默认作为文本文件读取 (.txt, .md 等)
        with open(path, encoding="utf-8") as f:
            return f.read()